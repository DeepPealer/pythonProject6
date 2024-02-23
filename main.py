# -*- coding: utf-8 -*-
import telebot
from telebot import types
from dateutil.relativedelta import relativedelta
from datetime import datetime
import json
import inflect
import time
import threading


TOKEN = '6772054408:AAGqtH_omTRL0011a4r4sZW2GVnpbeom7eE'
bot = telebot.TeleBot(TOKEN)


villa_list = {
    "Вилла 1": {
        "villa_name": "Вилла 1",
        "villa_location": "Локация 1",
        "villa_google_name": "Гугл карты 1"
    }
}


villa_name = {}
villa_location = {}
villa_google_name = {}

duration_to_time = {
    '3_months': relativedelta(months=3),
    '6_months': relativedelta(months=6),
    '1_year': relativedelta(years=1)
}
first_side_name = {}
first_side_nib = {}
first_side_address = {}
first_side_passport = {}
first_side_citizenship = {}
first_side_phone = {}
first_side_director = {}

second_side_name = {}
second_side_nib = {}
second_side_address = {}
second_side_passport = {}
second_side_citizenship = {}
second_side_phone = {}
second_side_director = {}

services = ['Электричество Безлимитно', 'Электричество до 1 млн рупий', 'Электричество до 500 тыс рупий', 'Электричество другое', 'Уборка 1р в неделю','Уборка 2р в неделю', 'Уборка 1р в 2 недели', 'Уборка другое',
            'Интернет 50 мбит', 'Интернет 100 мбит', 'Интернет 200 мбит', 'Бассейн 1р в неделю', 'Бассейн 2р в неделю', 'Бассейн 1р в 2 недели', 'Смена белья 1р в неделю', 'Смена белья 2р в неделю', 'Смена белья 1р в 2 недели', 'Субаренда разрешена']
selected_services = {}
service_descriptions = {
    'Электричество Безлимитно': 'Electricity',
    'Электричество до 1 млн рупий': 'Electricity up to 1 000 000 rupees per month',
    'Электричество до 500 тыс рупий': 'Electricity up to 500 000 rupees per month',
    'Электричество другое': '',
    'Уборка 1р в неделю': 'Cleaning 1 time a week',
    'Уборка 2р в неделю': 'Cleaning 2 times a week',
    'Уборка 1р в 2 недели': 'Cleaning 1 time every 2 weeks',
    'Уборка другое': '',
    'Интернет 50 мбит': 'Wi-fi 50 mbit',
    'Интернет 100 мбит': 'Wi-fi 100 mbit',
    'Интернет 200 мбит': 'Wi-fi 200 mbit',
    'Бассейн 1р в неделю': 'Pool cleaning 1 time a week',
    'Бассейн 2р в неделю': 'Pool cleaning 2 times a week',
    'Бассейн 1р в 2 недели': 'Pool cleaning 1 time every 2 weeks',
    'Смена белья 1р в неделю': 'Change of bed linen 1 time a week',
    'Смена белья 2р в неделю': 'Change of bed linen 2 times a week',
    'Смена белья 1р в 2 недели': 'Change of bed linen 1 time every 2 weeks',
    'Субаренда разрешена' : ''
}
isAuto = {}
isElectricity = {}
electricity_cost = {}
isCleaning = {}
cleaning_times = {}
isInternet = {}
internet_speed = {}
isPoll = {}
pool_times = {}
isChanging = {}
changing_times = {}
isSubArend = {}

state_history = {}
user_type = {}
first_side = {}
second_side = {}

natural_name = {}
natural_citizenship = {}
natural_passport = {}
natural_phone = {}
natural_address = {}

legal_name = {}
legal_director_name = {}
legal_director_citizenship = {}
legal_director_passport = {}
legal_director_passport_photo = {}
legal_nib = {}
legal_address = {}

contract_sides = {}
contract_start_date = {}
contract_duration = {}
villa_location = {}
villa_name = {}
commission_size = {}


rent_period = {}
rent_sum = {}
renewal_payment_period = {}
deposit_size = {}
included_services = {}




@bot.message_handler(commands=['start'])
def start(message):

    user_id = message.chat.id
    electricity_cost[user_id] = 0
    selected_services[user_id] = []

    if user_id in selected_services:
        selected_services[user_id].clear()
    isAuto[user_id] = False
    first_side[user_id] = ''
    second_side[user_id] = ''
    isCleaning[user_id] = False
    cleaning_times[user_id] = 0
    isInternet[user_id] = False
    internet_speed[user_id] = 0
    isPoll[user_id] = False
    pool_times[user_id] = 0
    isChanging[user_id] = False
    changing_times[user_id] = 0
    isSubArend[user_id] = False
    keyboard = types.InlineKeyboardMarkup()
    button1 = types.InlineKeyboardButton('Договор на управление виллой',
                                         callback_data='management')
    button2 = types.InlineKeyboardButton('Договор на помесячную аренду',
                                         callback_data='monthly')
    button3 = types.InlineKeyboardButton('Договор на долгосрочную аренду',
                                         callback_data='long_term')
    keyboard.add(button1)
    keyboard.add(button2)
    keyboard.add(button3)
    bot.send_message(message.chat.id,
                     f"Выберите тип договора",
                     reply_markup=keyboard)


@bot.message_handler(commands=['cancel'])
def cancel(message):
    user_id = message.chat.id
    bot.send_message(user_id, "Ввод данных отменен, введите команду /start чтобы начать заново")


@bot.callback_query_handler(func=lambda call: True)
def callback_handler(call):
    user_id = call.message.chat.id
    if user_id not in state_history:
        state_history[user_id] = []
    if call.data == 'management':
        user_type[user_id] = 'management'
        handle_fsides(call)
    if call.data == 'delete':
        ask_for_delete(call)
    if call.data == 'delete_villa':
        ask_for_delete_villa(call)
    if call.data == 'monthly':
        user_type[user_id] = 'monthly'
        handle_fsides(call)
    if call.data == 'long_term':
        user_type[user_id] = 'long_term'
        handle_fsides(call)

    if call.data == 'f_first_company':
        first_side[user_id] = 'first_company'
        first_side_name[user_id] = 'PT SABA REAL ESTATE'
        first_side_nib[user_id] = '1705230076844 '
        first_side_address[user_id] = 'Jl Raya Munggu-Tanah Lot No.888x, Desa/Kelurahan Cemagi, Kec. Mengwi, Kab. Badung, Provinsi Bali'
        handle_ssides(call)
    if call.data == 'f_second_company':
        first_side[user_id] = 'second_company'
        first_side_name[user_id] = 'PT SABA REAL ESTATE'
        first_side_nib[user_id] = '1705230076844 '
        first_side_address[
            user_id] = 'Jl Raya Munggu-Tanah Lot No.888x, Desa/Kelurahan Cemagi, Kec. Mengwi, Kab. Badung, Provinsi Bali'
        handle_ssides(call)
    if call.data == 'f_third_company':
        first_side[user_id] = 'third_company'
        first_side_name[user_id] = 'PT SABA REAL ESTATE'
        first_side_nib[user_id] = '1705230076844 '
        first_side_address[user_id] = 'Jl Raya Munggu-Tanah Lot No.888x, Desa/Kelurahan Cemagi, Kec. Mengwi, Kab. Badung, Provinsi Bali'
        handle_ssides(call)
    if call.data == 'f_natural_person':
        first_side[user_id] = 'natural_person'
        handle_natural_person(call)
    if call.data == 'f_legal_entity':
        first_side[user_id] = 'legal_entity'
        handle_legal_entity(call)

    if call.data == 's_first_company':
        second_side[user_id] = 'first_company'
        second_side_name[user_id] = 'PT SABA REAL ESTATE'
        second_side_nib[user_id] = '1705230076844 '
        second_side_address[
            user_id] = 'Jl Raya Munggu-Tanah Lot No.888x, Desa/Kelurahan Cemagi, Kec. Mengwi, Kab. Badung, Provinsi Bali'
        bot.delete_message(call.message.chat.id, call.message.id)
        if user_type[user_id] == 'management':
            handle_management(call)
        elif user_type[user_id] == 'monthly':
            handle_monthly(call)
        elif user_type[user_id] == 'long_term':
            handle_long_term(call)
    if call.data == 's_second_company':
        second_side[user_id] = 'second_company'
        second_side_name[user_id] = 'PT SABA REAL ESTATE'
        second_side_nib[user_id] = '1705230076844 '
        second_side_address[
            user_id] = 'Jl Raya Munggu-Tanah Lot No.888x, Desa/Kelurahan Cemagi, Kec. Mengwi, Kab. Badung, Provinsi Bali'
        bot.delete_message(call.message.chat.id, call.message.id)
        if user_type[user_id] == 'management':
            handle_management(call)
        elif user_type[user_id] == 'monthly':
            handle_monthly(call)
        elif user_type[user_id] == 'long_term':
            handle_long_term(call)
    if call.data == 's_third_company':
        second_side[user_id] = 'third_company'
        second_side_name[user_id] = 'PT SABA REAL ESTATE'
        second_side_nib[user_id] = '1705230076844 '
        second_side_address[
            user_id] = 'Jl Raya Munggu-Tanah Lot No.888x, Desa/Kelurahan Cemagi, Kec. Mengwi, Kab. Badung, Provinsi Bali'
        bot.delete_message(call.message.chat.id, call.message.id)
        if user_type[user_id] == 'management':
            handle_management(call)
        elif user_type[user_id] == 'monthly':
            handle_monthly(call)
        elif user_type[user_id] == 'long_term':
            handle_long_term(call)
    if call.data == 's_natural_person':
        second_side[user_id] = 'natural_person'
        bot.delete_message(call.message.chat.id, call.message.id)
        handle_natural_person(call)
    if call.data == 's_legal_entity':
        second_side[user_id] = 'legal_entity'
        bot.delete_message(call.message.chat.id, call.message.id)
        handle_legal_entity(call)



    if call.data == 'legal_repeat':
        handle_legal_entity(call)
    if call.data == 'legal_next':
        if first_side[user_id] == 'legal_entity':
            first_side_name[user_id] = legal_name[user_id]
            first_side_passport[user_id] = legal_director_passport[user_id]
            first_side_citizenship[user_id] = legal_director_citizenship[user_id]
            first_side_director[user_id] = legal_director_name[user_id]
            first_side_nib[user_id] = legal_nib[user_id]
            first_side_address[user_id] = legal_address[user_id]
            handle_ssides(call)
        else:
            second_side_name[user_id] = legal_name[user_id]
            second_side_passport[user_id] = legal_director_passport[user_id]
            second_side_citizenship[user_id] = legal_director_citizenship[user_id]
            second_side_director[user_id] = legal_director_name[user_id]
            second_side_nib[user_id] = legal_nib[user_id]
            second_side_address[user_id] = legal_address[user_id]
            if user_type[user_id] == 'management':
                handle_management(call)
            elif user_type[user_id] == 'monthly':
                handle_monthly(call)
            elif user_type[user_id] == 'long_term':
                handle_long_term(call)
    if call.data == 'villas_repeat':
        add_villa(call)
    if call.data == 'villas_next':
        callback_villa_next(call)
    if call.data == 'natural_repeat':
        handle_natural_person(call)

    if call.data == 'natural_next':
        if first_side[user_id] == 'natural_person':
            first_side_name[user_id] = natural_name[user_id]
            first_side_citizenship[user_id] = natural_citizenship[user_id]
            first_side_phone[user_id] = natural_phone[user_id]
            first_side_passport[user_id] = natural_passport[user_id]
            first_side_address[user_id] = natural_address[user_id]
            handle_ssides(call)
        else:
            second_side_name[user_id] = natural_name[user_id]
            second_side_citizenship[user_id] = natural_citizenship[user_id]
            second_side_phone[user_id] = natural_phone[user_id]
            second_side_passport[user_id] = natural_passport[user_id]
            second_side_address[user_id] = natural_address[user_id]
            if user_type[user_id] == 'management':
                handle_management(call)
            elif user_type[user_id] == 'monthly':
                handle_monthly(call)
            elif user_type[user_id] == 'long_term':
                handle_long_term(call)
    if call.data.startswith('villa_'):
        bot.delete_message(call.message.chat.id, call.message.id)
        callback_villa(call)
    if call.data == 'add_villa':
        add_villa(call)
    if call.data.startswith('first_'):
        data = load_data(user_id)
        entity = call.data[6:]
        first_side[user_id] = data[entity]['type']
        if first_side[user_id] == 'natural_person':
            first_side_name[user_id] = data[entity]['name']
            first_side_citizenship[user_id] = data[entity]['citizenship']
            first_side_phone[user_id] = data[entity]['phone']
            first_side_passport[user_id] = data[entity]['passport']
            first_side_address[user_id] = data[entity]['address']

            natural_name[user_id] = data[entity]['name']
            natural_passport[user_id] = data[entity]['passport']

        elif first_side[user_id] == 'legal_entity':
            first_side_name[user_id] = data[entity]['name']
            first_side_passport[user_id] = data[entity]['director_passport']
            first_side_citizenship[user_id] = data[entity]['director_citizenship']
            first_side_director[user_id] = data[entity]['director_name']
            first_side_nib[user_id] = data[entity]['nib']
            first_side_address[user_id] = data[entity]['address']

            legal_name[user_id] = data[entity]['name']
            legal_director_name[user_id] = data[entity]['director_name']
            legal_director_citizenship[user_id] = data[entity]['director_citizenship']
            legal_director_passport[user_id] = data[entity]['director_passport']
            legal_nib[user_id] = data[entity]['nib']
            legal_address[user_id] = data[entity]['address']
        handle_ssides(call)



    if call.data.startswith('second_'):
        data = load_data(user_id)
        entity = call.data[7:]
        second_side[user_id] = data[entity]['type']
        bot.delete_message(call.message.chat.id, call.message.id)
        if second_side[user_id] == 'natural_person':
            second_side_name[user_id] = data[entity]['name']
            second_side_passport[user_id] = data[entity]['passport']
            second_side_address[user_id] = data[entity]['address']
            second_side_citizenship[user_id] = data[entity]['citizenship']
            second_side_phone[user_id] = data[entity]['phone']

            natural_name[user_id] = data[entity]['name']
            natural_passport[user_id] = data[entity]['passport']
        elif second_side[user_id] == 'legal_entity':
            second_side_name[user_id] = data[entity]['name']
            second_side_director[user_id] = data[entity]['director_name']
            second_side_passport[user_id] = data[entity]['director_passport']
            second_side_citizenship[user_id] = data[entity]['director_citizenship']
            second_side_nib[user_id] = data[entity]['nib']
            second_side_address[user_id] = data[entity]['address']

            legal_name[user_id] = data[entity]['name']
            legal_director_name[user_id] = data[entity]['director_name']
            legal_director_citizenship[user_id] = data[entity]['director_citizenship']
            legal_director_passport[user_id] = data[entity]['director_passport']
            legal_nib[user_id] = data[entity]['nib']
            legal_address[user_id] = data[entity]['address']
        if user_type[user_id] == 'management':
            handle_management(call)
        elif user_type[user_id] == 'monthly':
            handle_monthly(call)
        elif user_type[user_id] == 'long_term':
            handle_long_term(call)


    if call.data == '3_months':
        contract_duration[user_id] = call.data
        bot.edit_message_text(chat_id=call.message.chat.id, message_id=call.message.message_id,
                              text='Включить автопродление?',
                              disable_web_page_preview=True)
        keyboard = types.InlineKeyboardMarkup()
        button1 = types.InlineKeyboardButton("Да", callback_data='auto_yes')
        button2 = types.InlineKeyboardButton("Нет", callback_data='auto_no')
        keyboard.add(button1, button2)
        bot.edit_message_reply_markup(chat_id=call.message.chat.id, message_id=call.message.message_id,
                                      reply_markup=keyboard)

    if call.data == '6_months':
        contract_duration[user_id] = call.data
        bot.edit_message_text(chat_id=call.message.chat.id, message_id=call.message.message_id,
                              text='Включить автопродление?',
                              disable_web_page_preview=True)
        keyboard = types.InlineKeyboardMarkup()
        button1 = types.InlineKeyboardButton("Да", callback_data='auto_yes')
        button2 = types.InlineKeyboardButton("Нет", callback_data='auto_no')
        keyboard.add(button1, button2)
        bot.edit_message_reply_markup(chat_id=call.message.chat.id, message_id=call.message.message_id,
                                      reply_markup=keyboard)

    if call.data == '1_year':
        contract_duration[user_id] = call.data
        bot.edit_message_text(chat_id=call.message.chat.id, message_id=call.message.message_id,
                              text='Включить автопродление?',
                              disable_web_page_preview=True)
        keyboard = types.InlineKeyboardMarkup()
        button1 = types.InlineKeyboardButton("Да", callback_data='auto_yes')
        button2 = types.InlineKeyboardButton("Нет", callback_data='auto_no')
        keyboard.add(button1, button2)
        bot.edit_message_reply_markup(chat_id=call.message.chat.id, message_id=call.message.message_id,
                                      reply_markup=keyboard)

    if call.data == 'auto_yes':
        isAuto[user_id] = True
        bot.delete_message(call.message.chat.id, call.message.id)
        process_contract_duration_step(call.message)

    if call.data == 'auto_no':
        isAuto[user_id] = False
        bot.delete_message(call.message.chat.id, call.message.id)
        process_contract_duration_step(call.message)

    if call.data == 'service_Электричество':
        bot.edit_message_text(chat_id=call.message.chat.id, message_id=call.message.message_id,
                              text='Какой тариф электричества?',
                              disable_web_page_preview=True)
        keyboard = types.InlineKeyboardMarkup()
        button1 = types.InlineKeyboardButton("Безлимитно", callback_data='elec_unl')
        button2 = types.InlineKeyboardButton("До 1 млн рупий", callback_data='elec_1mln')
        button3 = types.InlineKeyboardButton("До 500 тыс рупий", callback_data='elec500t')
        button4 = types.InlineKeyboardButton("Другая сумма", callback_data='elec_other')
        button5 = types.InlineKeyboardButton("Отмена", callback_data='elec_cancel')
        keyboard.add(button1, button2)
        keyboard.add(button3, button4)
        keyboard.add(button5)
        bot.edit_message_reply_markup(chat_id=call.message.chat.id, message_id=call.message.message_id,
                                      reply_markup=keyboard)

    if call.data == 'elec_cancel':
        isElectricity[user_id] = False
        electricity_cost[user_id] = 0
        process_deposit_size_step(call.message)

    if call.data == 'elec_unl':
        isElectricity[user_id] = True
        electricity_cost[user_id] = 'Безлимитно'
        process_deposit_size_step(call.message)

    elif call.data == 'elec_1mln':
        isElectricity[user_id] = True
        electricity_cost[user_id]= 'До 1 млн рупий'
        process_deposit_size_step(call.message)

    elif call.data == 'elec500t':
        isElectricity[user_id] = True
        electricity_cost[user_id] = 'До 500 тыс рупий'
        process_deposit_size_step(call.message)

    elif call.data == 'elec_other':
        msg = bot.send_message(chat_id=call.message.chat.id, text='Введите сумму:')
        bot.register_next_step_handler(msg, handle_user_response_electricity)


    if call.data == 'service_Уборка':
        bot.edit_message_text(chat_id=call.message.chat.id, message_id=call.message.message_id,
                              text='Как часто уборка?',
                              disable_web_page_preview=True)
        keyboard = types.InlineKeyboardMarkup()
        button1 = types.InlineKeyboardButton("1 р в неделю", callback_data='clean_one')
        button2 = types.InlineKeyboardButton("2 р в неделю", callback_data='clean_two')
        button3 = types.InlineKeyboardButton("3 р в неделю", callback_data='clean_three')
        button4 = types.InlineKeyboardButton("Другое", callback_data='clean_other')
        button5 = types.InlineKeyboardButton("Отмена", callback_data='clean_cancel')
        keyboard.add(button1)
        keyboard.add(button2)
        keyboard.add(button3)
        keyboard.add(button4)
        keyboard.add(button5)
        bot.edit_message_reply_markup(chat_id=call.message.chat.id, message_id=call.message.message_id,
                                      reply_markup=keyboard)

    if call.data == 'clean_one':
        isCleaning[user_id] = True
        cleaning_times[user_id] = '3 р в неделю'
        process_deposit_size_step(call.message)

    if call.data == 'clean_two':
        isCleaning[user_id] = True
        cleaning_times[user_id] = '2 р в неделю'
        process_deposit_size_step(call.message)

    elif call.data == 'clean_three':
        isCleaning[user_id] = True
        cleaning_times[user_id]= '3 р в неделю'
        process_deposit_size_step(call.message)

    elif call.data == 'clean_other':
        msg = bot.send_message(chat_id=call.message.chat.id, text='Введите количество раз:')
        bot.register_next_step_handler(msg, handle_user_response_cleaning)

    if call.data == 'clean_cancel':
        isCleaning[user_id] = False
        cleaning_times[user_id] = 0
        process_deposit_size_step(call.message)


    if call.data == 'service_Интернет':
        bot.edit_message_text(chat_id=call.message.chat.id, message_id=call.message.message_id,
                              text='Какая скорость интернета?',
                              disable_web_page_preview=True)
        keyboard = types.InlineKeyboardMarkup()
        button1 = types.InlineKeyboardButton("50 мбит", callback_data='internet_50')
        button2 = types.InlineKeyboardButton("100 мбит", callback_data='internet_100')
        button3 = types.InlineKeyboardButton("200 мбит", callback_data='internet_200')
        button5 = types.InlineKeyboardButton("Отмена", callback_data='internet_cancel')
        keyboard.add(button1)
        keyboard.add(button2)
        keyboard.add(button3)
        keyboard.add(button5)
        bot.edit_message_reply_markup(chat_id=call.message.chat.id, message_id=call.message.message_id,
                                      reply_markup=keyboard)

    if call.data == 'internet_50':
        isInternet[user_id] = True
        internet_speed[user_id] = '50 мбит'
        process_deposit_size_step(call.message)

    if call.data == 'internet_100':
        isInternet[user_id] = True
        internet_speed[user_id] = '100 мбит'
        process_deposit_size_step(call.message)

    elif call.data == 'internet_200':
        isInternet[user_id] = True
        internet_speed[user_id] = '200 мбит'
        process_deposit_size_step(call.message)

    if call.data == 'internet_cancel':
        isInternet[user_id] = False
        internet_speed[user_id] = 0
        process_deposit_size_step(call.message)


    if call.data == 'service_Обслуживание бассейна':
        bot.edit_message_text(chat_id=call.message.chat.id, message_id=call.message.message_id,
                              text='Сколько раз в неделю обслуживают бассейн?',
                              disable_web_page_preview=True)
        keyboard = types.InlineKeyboardMarkup()
        button1 = types.InlineKeyboardButton("1р в неделю", callback_data='pool_1tw')
        button2 = types.InlineKeyboardButton("2р в неделю", callback_data='pool_2tw')
        button3 = types.InlineKeyboardButton("1р в две недели", callback_data='pool_1t2w')
        button5 = types.InlineKeyboardButton("Отмена", callback_data='pool_cancel')
        keyboard.add(button1)
        keyboard.add(button2)
        keyboard.add(button3)
        keyboard.add(button5)
        bot.edit_message_reply_markup(chat_id=call.message.chat.id, message_id=call.message.message_id,
                                      reply_markup=keyboard)

    if call.data == 'pool_1tw':
        isPoll[user_id] = True
        pool_times[user_id] = '1р в неделю'
        process_deposit_size_step(call.message)

    if call.data == 'pool_2tw':
        isPoll[user_id] = True
        pool_times[user_id] = '2р в неделю'
        process_deposit_size_step(call.message)

    elif call.data == 'pool_1t2w':
        isPoll[user_id] = True
        pool_times[user_id] = '1р в две недели'
        process_deposit_size_step(call.message)

    if call.data == 'pool_cancel':
        isPoll[user_id] = False
        pool_times[user_id] = 0
        process_deposit_size_step(call.message)


    if call.data == 'service_Смена постельного белья':
        bot.edit_message_text(chat_id=call.message.chat.id, message_id=call.message.message_id,
                              text='Сколько раз в неделю меняют постельное бельё?',
                              disable_web_page_preview=True)
        keyboard = types.InlineKeyboardMarkup()
        button1 = types.InlineKeyboardButton("1р в неделю", callback_data='change_1tw')
        button2 = types.InlineKeyboardButton("2р в неделю", callback_data='change_2tw')
        button3 = types.InlineKeyboardButton("1р в две недели", callback_data='change_1t2w')
        button5 = types.InlineKeyboardButton("Отмена", callback_data='change_cancel')
        keyboard.add(button1)
        keyboard.add(button2)
        keyboard.add(button3)
        keyboard.add(button5)
        bot.edit_message_reply_markup(chat_id=call.message.chat.id, message_id=call.message.message_id,
                                      reply_markup=keyboard)

    if call.data == 'change_1tw':
        isChanging[user_id] = True
        changing_times[user_id] = '1р в неделю'
        process_deposit_size_step(call.message)

    if call.data == 'change_2tw':
        isChanging[user_id] = True
        changing_times[user_id] = '2р в неделю'
        process_deposit_size_step(call.message)

    elif call.data == 'change_1t2w':
        isChanging[user_id] = True
        changing_times[user_id] = '1р в две недели'
        process_deposit_size_step(call.message)

    if call.data == 'change_cancel':
        isChanging[user_id] = False
        changing_times[user_id] = 0
        process_deposit_size_step(call.message)


    if call.data == 'service_Разрешена субаренда':
        bot.edit_message_text(chat_id=call.message.chat.id, message_id=call.message.message_id,
                              text='Субаренда разрешена?',
                              disable_web_page_preview=True)
        keyboard = types.InlineKeyboardMarkup()
        button1 = types.InlineKeyboardButton("Да", callback_data='sub_yes')
        button2 = types.InlineKeyboardButton("Нет", callback_data='sub_no')
        button5 = types.InlineKeyboardButton("Отмена", callback_data='sub_cancel')
        keyboard.add(button1)
        keyboard.add(button2)
        keyboard.add(button5)
        bot.edit_message_reply_markup(chat_id=call.message.chat.id, message_id=call.message.message_id,
                                      reply_markup=keyboard)

    if call.data == 'sub_yes':
        isChanging[user_id] = True
        process_deposit_size_step(call.message)

    if call.data == 'sub_no':
        isChanging[user_id] = False
        process_deposit_size_step(call.message)

    if call.data == 'sub_cancel':
        process_deposit_size_step(call.message)

    if call.data == 'ok':
        bot.edit_message_text(chat_id=call.message.chat.id, message_id=call.message.message_id,
                              text='Идет создание документа...',
                              disable_web_page_preview=True)
        create_agreement(call.message)
        start(call.message)
    if call.data == 'natural_save':
        natural_save(call)
    if call.data == 'legal_save':
        legal_save(call)
    if call.data == 'villas_save':
        villa_save(call)
    elif call.data.startswith('service_'):
        if call.data == 'service_Электричество другое':
            msg = bot.send_message(chat_id=call.message.chat.id, text='Введите сумму:')
            bot.register_next_step_handler(msg, handle_user_response_electricity)
            selected_services.setdefault(user_id, []).append('Электричество другое')

        elif call.data == 'service_Уборка другое':
            msg = bot.send_message(chat_id=call.message.chat.id, text='Введите количество раз в неделю:')
            bot.register_next_step_handler(msg, handle_user_response_cleaning)
            selected_services.setdefault(user_id, []).append('Уборка другое')

        else:
            service = call.data.replace('service_', '')
            if service in selected_services.get(user_id, []):
                selected_services[user_id].remove(service)
            else:
                selected_services.setdefault(user_id, []).append(service)
            bot.edit_message_text(chat_id=call.message.chat.id, message_id=call.message.message_id,
                                  text='Что включено?',
                                  reply_markup=generate_keyboard_requestions(user_id))


def create_agreement(message):
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.shared import Cm

    user_id = message.chat.id
    villa_list_1 = load_data_villa(user_id)
    villa_list.update(villa_list_1)
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(2.36)
    section.right_margin = Cm(2.4)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    style = doc.styles['Normal']
    style.paragraph_format.space_after = Cm(0)

    # Создание стиля для текста шрифтом Times New Roman размером 11
    times_new_roman_11_style = doc.styles['Normal']
    times_new_roman_11_style.font.name = 'Times New Roman'
    times_new_roman_11_style.font.size = Pt(11)


    center_bold_style = doc.styles.add_style('CenterBold', WD_PARAGRAPH_ALIGNMENT.CENTER)
    center_bold_style.font.bold = True
    if user_type[user_id] == 'management':
        paragraph = doc.add_paragraph(style=center_bold_style)
        run = paragraph.add_run('PROPERTY MANAGEMENT AGREEMENT')
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_paragraph(f'This Management Contract ("Contract") is entered into on this {contract_start_date[user_id]} between the Owner(s) '
                          'of the Villa, hereinafter referred to as the "Owner," and the Management Company, hereinafter '
                          'referred to as the "Manager."')
        doc.add_paragraph()

        # Добавление контрактных данных
        doc.add_paragraph('This Management contract is made by :\n', style=times_new_roman_11_style)

        match first_side[user_id]:
            case 'first_company':
                doc.add_paragraph(
                    '1. Name 	 : PT SABA REAL ESTATE\n'
                    '    NIB 		 : 1705230076844\n'
                    '    Address  	 : Jl Raya Munggu-Tanah Lot No.888x, Desa/Kelurahan Cemagi, Kec. Mengwi, Kab. Badung, Provinsi Bali \n'
                    f'    Director   	 : \n',
                    style=times_new_roman_11_style)
            case 'second_company':
                doc.add_paragraph(
                    '1. Name 	 : PT SABA REAL ESTATE\n'
                    '    NIB 		 : 1705230076844\n'
                    '    Address  	 : Jl Raya Munggu-Tanah Lot No.888x, Desa/Kelurahan Cemagi, Kec. Mengwi, Kab. Badung, Provinsi Bali \n'
                    f'    Director   	 : \n',
                    style=times_new_roman_11_style)
            case 'third_company':
                doc.add_paragraph(
                    '1. Name 	 : PT SABA REAL ESTATE\n'
                    '    NIB 		 : 1705230076844\n'
                    '    Address  	 : Jl Raya Munggu-Tanah Lot No.888x, Desa/Kelurahan Cemagi, Kec. Mengwi, Kab. Badung, Provinsi Bali \n'
                    f'    Director   	 : \n',
                    style=times_new_roman_11_style)
            case 'natural_person':
                doc.add_paragraph(
                    f'1. Name 	 : {first_side_name[user_id]}\n'
                    f'    Passport 	 : {first_side_passport[user_id]}\n'
                    f'    Citizen           : {first_side_citizenship[user_id]}\n'
                    f'    Phone number:{first_side_phone[user_id]}\n'
                    f'    Address          :{first_side_address[user_id]}\n',
                    style=times_new_roman_11_style)
            case 'legal_entity':
                doc.add_paragraph(
                    f'1. Name 	 : {first_side_name[user_id]}\n'
                    f'    NIB 		 : {first_side_nib[user_id]}\n'
                    f'    Address  	 : {first_side_address[user_id]}\n'
                    f'    Director       : {first_side_director[user_id]}, {first_side_citizenship[user_id]}, {first_side_passport[user_id]}\n',
                    style=times_new_roman_11_style)


        doc.add_paragraph(' To be known as Manager\n\n', style=times_new_roman_11_style)
        match second_side[user_id]:
            case 'first_company':
                doc.add_paragraph(
                    '2. Name 	 : PT SABA REAL ESTATE\n'
                    '    NIB 		 : 1705230076844\n'
                    '    Address  	 : Jl Raya Munggu-Tanah Lot No.888x, Desa/Kelurahan Cemagi, Kec. Mengwi, Kab. Badung, Provinsi Bali \n',
                    style=times_new_roman_11_style)
            case 'second_company':
                doc.add_paragraph(
                    '2. Name 	 : PT SABA REAL ESTATE\n'
                    '    NIB 		 : 1705230076844\n'
                    '    Address  	 : Jl Raya Munggu-Tanah Lot No.888x, Desa/Kelurahan Cemagi, Kec. Mengwi, Kab. Badung, Provinsi Bali \n',
                    style=times_new_roman_11_style)
            case 'third_company':
                doc.add_paragraph(
                    f'2. Name 	 : PT SABA REAL ESTATE\n'
                    f'    NIB 		 : 1705230076844\n'
                    '    Address  	 : Jl Raya Munggu-Tanah Lot No.888x, Desa/Kelurahan Cemagi, Kec. Mengwi, Kab. Badung, Provinsi Bali \n',
                    style=times_new_roman_11_style)
            case 'natural_person':
                doc.add_paragraph(
                    f'2. Name 	 : {second_side_name[user_id]}\n'
                    f'    Passport 	 : {second_side_passport[user_id]}\n'
                    f'    Citizen           : {second_side_citizenship[user_id]}\n'
                    f'    Phone number:{second_side_phone[user_id]}\n'
                    f'    Address          :{second_side_address[user_id]}\n',
                    style=times_new_roman_11_style)
            case 'legal_entity':
                doc.add_paragraph(
                    f'2. Name 	 : {second_side_name[user_id]}\n'
                    f'    NIB 		 : {second_side_nib[user_id]}\n'
                    f'    Address  	 : {second_side_address[user_id]}\n'
                    f'    Director       : {second_side_director[user_id]}, {first_side_citizenship[user_id]}, {first_side_passport[user_id]}\n',
                    style=times_new_roman_11_style)

        doc.add_paragraph(' To be known as Owner\n\n', style=times_new_roman_11_style)



        doc.add_paragraph()


        doc.add_paragraph('1. PROPERTY DETAILS:', style=times_new_roman_11_style)

        doc.add_paragraph(f'   1.1 The Owner owns the property located at {villa_location[user_id]}, hereinafter'
                          ' referred to as the "Villa."', style=times_new_roman_11_style)



        doc.add_paragraph()


        doc.add_paragraph('2. APPOINTMENT OF MANAGER:', style=times_new_roman_11_style)

        rent_period_time = duration_to_time[contract_duration[user_id]]
        start_date = datetime.strptime(contract_start_date[user_id], '%d.%m.%Y')
        end_date = start_date + rent_period_time
        end_date_str = end_date.strftime('%d.%m.%Y')

        doc.add_paragraph('   2.1 The Owner appoints the Manager as the exclusive manager of the Villa'
                          f' for a period of {contract_duration[user_id].replace("_", " ")}, commencing on {contract_start_date[user_id]} and ending on {end_date_str}.',
                          style=times_new_roman_11_style)
        if isAuto[user_id]:
            doc.add_paragraph('   If neither party notifies the other of the termination of cooperation at least 3 months in advance, the contract is automatically prolonged',
                              style=times_new_roman_11_style)

        doc.add_paragraph('   2.2 The Manager accepts the appointment and agrees to manage the Villa on behalf of the Owner during the term of this Contract.',
                          style=times_new_roman_11_style)


        doc.add_paragraph()


        doc.add_paragraph('3. SERVICES PROVIDED BY MANAGER:', style=times_new_roman_11_style)

        doc.add_paragraph('   3.1 The Manager shall provide the following services:', style=times_new_roman_11_style)

        doc.add_paragraph('       - Marketing and advertising the Villa for rental purposes.', style=times_new_roman_11_style)

        doc.add_paragraph('       - Booking management, including handling reservations, guest inquiries, and payments.', style=times_new_roman_11_style)

        doc.add_paragraph('       - Check-in and check-out procedures for guests.', style=times_new_roman_11_style)

        doc.add_paragraph('       - Coordination of cleaning, maintenance, and repairs.', style=times_new_roman_11_style)

        doc.add_paragraph('       - Guest support and concierge services.', style=times_new_roman_11_style)

        doc.add_paragraph('       - Accounting and financial reporting.', style=times_new_roman_11_style)

        doc.add_paragraph('   3.2 The Manager shall exercise reasonable care and diligence in performing the services.', style=times_new_roman_11_style)


        doc.add_paragraph()

        doc.add_paragraph('4. TERM AND TERMINATION:')

        doc.add_paragraph('   4.1 The initial term of this Contract shall be 1 year. Either party may terminate this Contract with 3 month written notice prior to the intended termination date.')

        doc.add_paragraph('   4.2 Termination shall not relieve either party of any obligations arising prior to the termination date.')

        doc.add_paragraph('   4.3 In the event of termination, the Manager shall return any funds, records, or documents belonging to the Owner within 10 days of termination.')

        doc.add_paragraph('   4.4 Property is considered transferred under management of the Company after confirmation by '
                          'the Owner that the property is ready to receive guests, the Owner provides good quality '
                          'property photos to the Company (or request Company to make it), transfers all OTA listings (if '
                          'there are any), fills out the checklist and gets confirmation from the Company that all necessary documents have been received.')


        doc.add_paragraph()


        doc.add_paragraph('5. COMPENSATION:')


        doc.add_paragraph(f'   5.1 The Owner shall pay the Manager a commission of {commission_size[user_id]} of the gross rental income generated from the Villa.')

        doc.add_paragraph('   5.2 The Manager shall provide a monthly statement detailing rental income, expenses, and the net amount payable to the Owner.')

        doc.add_paragraph('   5.3 The Manager shall disburse the net amount to the Owner within 10'
                        ' days after the end of each calendar month.')


        doc.add_paragraph()


        doc.add_paragraph('6. INSURANCE:')

        doc.add_paragraph('   6.1 The Manager shall not be liable for any loss, damage, or injury to'
                          ' the Villa, its contents, or any third parties, except in cases of'
                          ' negligence or willful misconduct.')


        doc.add_paragraph()


        doc.add_paragraph('7. GOVERNING LAW:')

        doc.add_paragraph('   7.1 This Contract shall be governed by and construed in accordance with'
                          ' the laws of Indonesia.\n')

        doc.add_paragraph('IN WITNESS WHEREOF, the parties hereto have executed this Pre-Management Contract as of the date first above written.\n')

        # Добавление таблицы
        table = doc.add_table(rows=1, cols=2)
        cell_1 = table.cell(0, 0)
        cell_1.text = f'Management\n\n\n\n\n\n\n{first_side_name[user_id]}'
        cell_2 = table.cell(0, 1)
        cell_2.text = f'Owner\n\n\n\n\n\n\n{second_side_name[user_id]}'
        for paragraph in cell_1.paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for paragraph in cell_2.paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        for paragraph in cell_1.paragraphs:
            for run in paragraph.runs:
                run.bold = True
        for paragraph in cell_2.paragraphs:
            for run in paragraph.runs:
                run.bold = True

        # Сохранение файла
        doc.save('Property_Management_Agreement.docx')
        bot.send_document(chat_id=user_id, document=open('Property_Management_Agreement.docx', 'rb'))
    elif user_type[user_id]  == 'monthly':
        paragraph = doc.add_paragraph(style=center_bold_style)
        run = paragraph.add_run('LEASE AGREEMENT')
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph('This Lease Agreement is made by: :\n', style=times_new_roman_11_style)
        match first_side[user_id]:
            case 'first_company':
                doc.add_paragraph(
                    '1. Name 	 : PT SABA REAL ESTATE\n'
                    '    NIB 		 : 1705230076844\n'
                    '    Address  	 : Jl Raya Munggu-Tanah Lot No.888x, Desa/Kelurahan Cemagi, Kec. Mengwi, Kab. Badung, Provinsi Bali \n'
                    f'    Director   	 : \n',
                    style=times_new_roman_11_style)
            case 'second_company':
                doc.add_paragraph(
                    '1. Name 	 : PT SABA REAL ESTATE\n'
                    '    NIB 		 : 1705230076844\n'
                    '    Address  	 : Jl Raya Munggu-Tanah Lot No.888x, Desa/Kelurahan Cemagi, Kec. Mengwi, Kab. Badung, Provinsi Bali \n'
                    f'    Director   :\n',
                    style=times_new_roman_11_style)
            case 'third_company':
                doc.add_paragraph(
                    f'1. Name 	 : PT SABA REAL ESTATE\n'
                    f'    NIB 		 : 1705230076844\n'
                    '    Address  	 : Jl Raya Munggu-Tanah Lot No.888x, Desa/Kelurahan Cemagi, Kec. Mengwi, Kab. Badung, Provinsi Bali \n'
                    f'    Director   : \n',
                    style=times_new_roman_11_style)
            case 'natural_person':
                doc.add_paragraph(
                    f'1. Name 	 : {first_side_name[user_id]}\n'
                    f'    Passport 	 : {first_side_passport[user_id]}\n'
                    f'    Citizen           : {first_side_citizenship[user_id]}\n'
                    f'    Phone number:{first_side_phone[user_id]}\n'
                    f'    Address          :{first_side_address[user_id]}\n',
                    style=times_new_roman_11_style)
            case 'legal_entity':
                doc.add_paragraph(
                    f'1. Name 	 : {first_side_name[user_id]}\n'
                    f'    NIB 		 : {first_side_nib[user_id]}\n'
                    f'    Address  	 : {first_side_address[user_id]}\n'
                    f'    Director       : {first_side_director[user_id]}, {first_side_citizenship[user_id]}, {first_side_passport[user_id]}\n',
                    style=times_new_roman_11_style)

        doc.add_paragraph(' To be known as Property Owner\n\n', style=times_new_roman_11_style)
        match second_side[user_id]:
            case 'first_company':
                doc.add_paragraph(
                    '2. Name 	 : PT SABA REAL ESTATE\n'
                    '    NIB 		 : 1705230076844\n'
                    '    Address  	 : Jl Raya Munggu-Tanah Lot No.888x, Desa/Kelurahan Cemagi, Kec. Mengwi, Kab. Badung, Provinsi Bali \n'
                    f'    Director   : \n',
                    style=times_new_roman_11_style)
            case 'second_company':
                doc.add_paragraph(
                    '2. Name 	 : PT SABA REAL ESTATE\n'
                    '    NIB 		 : 1705230076844\n'
                    '    Address  	 : Jl Raya Munggu-Tanah Lot No.888x, Desa/Kelurahan Cemagi, Kec. Mengwi, Kab. Badung, Provinsi Bali \n'
                    f'    Director   : \n',
                    style=times_new_roman_11_style)
            case 'third_company':
                doc.add_paragraph(
                    f'2. Name 	 : PT SABA REAL ESTATE\n'
                    f'    NIB 		 : 1705230076844\n'
                    '    Address  	 : Jl Raya Munggu-Tanah Lot No.888x, Desa/Kelurahan Cemagi, Kec. Mengwi, Kab. Badung, Provinsi Bali \n',
                    f'    Director          : \n',
                    style=times_new_roman_11_style)
            case 'natural_person':
                doc.add_paragraph(
                    f'2. Name 	 : {second_side_name[user_id]}\n'
                    f'    Passport 	 : {second_side_passport[user_id]}\n'
                    f'    Citizen           : {second_side_citizenship[user_id]}\n'
                    f'    Phone number:{second_side_phone[user_id]}\n'
                    f'    Address          :{second_side_address[user_id]}\n',
                    style=times_new_roman_11_style)
            case 'legal_entity':
                doc.add_paragraph(
                    f'2. Name 	 : {second_side_name[user_id]}\n'
                    f'    NIB 		 : {second_side_nib[user_id]}\n'
                    f'    Address  	 : {second_side_address[user_id]}\n'
                    f'    Director          : {second_side_director[user_id]}, {second_side_citizenship[user_id]}, {second_side_passport[user_id]}\n',
                    style=times_new_roman_11_style)

        doc.add_paragraph(' To be known as Tenant\n\n', style=times_new_roman_11_style)

        rent_period_time = str_to_relativedelta(rent_period[user_id])
        start_date = datetime.strptime(contract_start_date[user_id], '%d.%m.%Y')
        end_date = start_date + rent_period_time
        end_date_str = end_date.strftime('%d.%m.%Y')

        p = inflect.engine()

        number_with_spaces = '{:,}'.format(int(rent_sum[user_id])).replace(',', ' ')
        words = p.number_to_words(rent_sum[user_id])

        nws = '{:,}'.format(int(deposit_size[user_id])).replace(',', ' ')
        w = p.number_to_words(deposit_size[user_id])

        month_word = "month" if rent_period[user_id] == '1' else "months"
        doc.add_paragraph(f'Conditions:')
        doc.add_paragraph(f'Property owner hereby agrees to lease the property Located of {villa_list[chosen_villa]["villa_google_name"]}\n')

        num_months = rent_period[user_id].split('_')[0]

        doc.add_paragraph(f'The Lease Period start on {contract_start_date[user_id]} and ends on {end_date_str} ({num_months} {month_word})\n')
        doc.add_paragraph(f'The Agreed Rental Value per month of {number_with_spaces} ({words}) IDR \n')
        lease_string = "The Leased property will be used for residential purposes"
        if "Субаренда разрешена" in selected_services[user_id]:
            lease_string += ", but the Tenant can also sublease this villa"
        number_with_spaces = '{:,}'.format(int(electricity_cost[user_id])).replace(',', ' ')
        def get_service_description(key, user_id):
            if key == 'Электричество другое':
                return f'Electricity up to {number_with_spaces} rupees per month'
            elif key == 'Уборка другое':
                return f'Cleaning {cleaning_times[user_id]} times a week'
            else:
                return service_descriptions.get(key, '')

        included_services = [get_service_description(service, user_id) for service in selected_services[user_id] if
                             service != 'Субаренда разрешена']
        included_services_text = "The villa rental price also includes:\n- " + "\n- ".join(included_services)
        doc.add_paragraph(f'The following terms and conditions that must be agreed during the rental period are :\n\n1.{lease_string}\n\n2. {included_services_text} \n\n3. The leased property is equipped with furniture and a list of inventory / buildings that are known to the parties.\n\n4. The property owner can inspect the property any time after giving it notice 1 days before.\n\n5. Property is not allowed to carry out activities that violate national Law and Local Customary law.\n\n6. Loss or damage by tenant property is the responsibility tenants at a fee agreed by owner and tenants are prohibited hard to change all form of buildings that are rented out during the leasen period in the form whatever.\n\n7. Tenants is required at the end of this agreement to handback the Villa Building as well as the stove and water heater to their original state , namely conditions at the initial time rented by the tenants to the Owner in state or condition is good and empty without occupants.\n\n8.The Owner of the Property shall collect from the Tenant a security deposit in the amount of {nws} ({w}) rupees as security for the property.\nThe deposit is returned to the Tenant upon expiration of this Rental Agreement, upon acceptance of the villa by the manager and provided that an inspection of the premises does not reveal any damage. If such damage is discovered, the cost of repairing it will be deducted from the deposit amount, and the remainder of the deposit will be returned to the Tenant.\n\n9.The tenant agrees to pay for the next month {renewal_payment_period[user_id]} days before the start of the next month.\n\n10. If the tenant wants to terminate the contract early, the deposit is not refundable.\n\nProperty owner and tenants hereby acknowledge and understand the terms and conditions written in this agreement.\n\n')

        table = doc.add_table(rows=1, cols=2)
        cell_1 = table.cell(0, 0)
        cell_1.text = f'Property Owner\n\n\n\n\n\n\n{first_side_name[user_id]}'
        cell_2 = table.cell(0, 1)
        cell_2.text = f'Tenant\n\n\n\n\n\n\n{second_side_name[user_id]}'
        for paragraph in cell_1.paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for paragraph in cell_2.paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        for paragraph in cell_1.paragraphs:
            for run in paragraph.runs:
                run.bold = True
        for paragraph in cell_2.paragraphs:
            for run in paragraph.runs:
                run.bold = True

        doc.save('Lease_Agreement.docx')
        bot.send_document(chat_id=user_id, document=open('Lease_Agreement.docx', 'rb'))



def handle_user_response_electricity(message):
    if message.text.isdigit():
        user_id = message.chat.id
        isElectricity[user_id] = True
        electricity_cost[user_id] = message.text
        process_deposit_size_step(message)
    else:
        bot.send_message(chat_id=message.chat.id, text='Пожалуйста, введите числовое значение.')

def handle_user_response_cleaning(message):
    if message.text.isdigit():
        user_id = message.chat.id
        isCleaning[user_id] = True
        cleaning_times[user_id] = message.text
        process_deposit_size_step(message)
    else:
        bot.send_message(chat_id=message.chat.id, text='Пожалуйста, введите числовое значение.')


def handle_management(call):
    user_id = call.message.chat.id
    contract_sides[user_id] = call.message.text
    msg = bot.send_message(user_id, "Дата старта контракта(в формате дд.мм.гггг)")
    bot.register_next_step_handler(msg, process_contract_start_date_step)

def is_valid_date(date_text):
    try:
        datetime.strptime(date_text, '%d.%m.%Y')
        return True
    except ValueError:
        return False

def process_contract_start_date_step(message):
    user_id = message.chat.id
    date_text = message.text
    if not is_valid_date(date_text):
        msg = bot.send_message(user_id,
                               "Дата введена в неверном формате. Пожалуйста, введите дату в формате дд.мм.гггг")
        bot.register_next_step_handler(msg, process_contract_start_date_step)
        return
    else:
        contract_start_date[user_id] = message.text
    keyboard = types.InlineKeyboardMarkup()
    button1 = types.InlineKeyboardButton("3 мес", callback_data='3_months')
    button2 = types.InlineKeyboardButton("6 мес", callback_data='6_months')
    button3 = types.InlineKeyboardButton("1 год", callback_data='1_year')
    keyboard.add(button1, button2, button3)
    bot.send_message(user_id, "Длительность контракта", reply_markup=keyboard)


def process_contract_duration_step(message):
    user_id = message.chat.id
    msg = bot.send_message(user_id, "Локация виллы")
    bot.register_next_step_handler(msg, process_villa_location_step)


def process_villa_location_step(message):
    user_id = message.chat.id
    villa_location[user_id] = message.text
    msg = bot.send_message(user_id, "Название виллы")
    bot.register_next_step_handler(msg, process_villa_name_step)

def process_villa_name_step(message):
    user_id = message.chat.id
    villa_name[user_id] = message.text
    msg = bot.send_message(user_id, "Размер комиссии")
    bot.register_next_step_handler(msg, process_commission_size_step)

def process_commission_size_step(message):
    user_id = message.chat.id
    commission_size[user_id] = message.text
    bot.send_message(user_id, "Идет создание документа...")
    create_agreement(message)
    start(message)



def handle_monthly(call):
    bot.send_message(call.message.chat.id, "Пожалуйста, введите следующую информацию:")
    choose_villa(call.message)

def choose_villa(message):
    user_id = message.chat.id
    data = load_data_villa(user_id)

    keyboard = types.InlineKeyboardMarkup()
    for villa in data:
        button = types.InlineKeyboardButton(villa, callback_data=f'villa_{villa}')
        keyboard.add(button)
    button_add = types.InlineKeyboardButton("Ввести виллу", callback_data='add_villa')
    button_delete = types.InlineKeyboardButton("Удалить виллу", callback_data='delete_villa')
    keyboard.add(button_add)
    keyboard.add(button_delete)
    bot.send_message(message.chat.id, "Какая вилла", reply_markup=keyboard)

def callback_villa_next(call):
    user_id = call.message.chat.id
    global chosen_villa
    chosen_villa = villa_name[user_id]
    msg = bot.send_message(user_id, "Когда старт контракта(в формате дд.мм.гггг)")
    bot.register_next_step_handler(msg, process_contract_start_date)

def callback_villa(call):
    user_id = call.message.chat.id
    global chosen_villa
    chosen_villa = call.data[6:]

    msg = bot.send_message(user_id, "Когда старт контракта(в формате дд.мм.гггг)")
    bot.register_next_step_handler(msg, process_contract_start_date)


def process_contract_start_date(message):
    user_id = message.chat.id
    date_text = message.text
    if not is_valid_date(date_text):
        msg = bot.send_message(user_id,
                               "Дата введена в неверном формате. Пожалуйста, введите дату в формате дд.мм.гггг")
        bot.register_next_step_handler(msg, process_contract_start_date)
        return
    else:
        contract_start_date[user_id] = message.text
    msg = bot.send_message(user_id, "На какой период хотят арендовать, в месяцах")
    bot.register_next_step_handler(msg, process_rent_period_step)


def process_rent_period_step(message):
    user_id = message.chat.id
    rent_period_text = message.text

    if not rent_period_text.isdigit():
        msg = bot.send_message(user_id,
                               "Введенное значение не является числом. Пожалуйста, введите период аренды в месяцах.")
        bot.register_next_step_handler(msg, process_rent_period_step)
        return

    rent_period[user_id] = rent_period_text + '_month'
    msg = bot.send_message(user_id, "Какая сумма аренды")
    bot.register_next_step_handler(msg, process_rent_sum_step)



def process_rent_sum_step(message):
    user_id = message.chat.id
    rent_sum_text = message.text

    if not rent_sum_text.isdigit():
        msg = bot.send_message(user_id,
                               "Введенное значение не является числом. Пожалуйста, введите сумму аренды.")
        bot.register_next_step_handler(msg, process_rent_sum_step)
        return
    else:
        rent_sum[user_id] = message.text
    msg = bot.send_message(user_id, "Для продления аренды платить за …… дней")
    bot.register_next_step_handler(msg, process_renewal_payment_period_step)


def process_renewal_payment_period_step(message):
    user_id = message.chat.id
    renewal_payment_period[user_id] = message.text
    msg = bot.send_message(user_id, "Размер депозита")
    bot.register_next_step_handler(msg, process_deposit_size_step)


def process_deposit_size_step(message):
    user_id = message.chat.id
    deposit_size_text = message.text
    if not deposit_size_text.isdigit():
        msg = bot.send_message(user_id,
                               "Введенное значение не является числом. Пожалуйста, введите размер депозита.")
        bot.register_next_step_handler(msg, process_deposit_size_step)
        return

    deposit_size[user_id] = message.text
    keyboard = types.InlineKeyboardMarkup()
    #services = ["Электричество", "Уборка", "Интернет", "Обслуживание бассейна", "Смена постельного белья", "Разрешена субаренда"]
    #for service in services:
      #  button = types.InlineKeyboardButton(service, callback_data=f'service_{service}')
     #   keyboard.add(button)
   # button3 = types.InlineKeyboardButton("Готово", callback_data='ok')
    #keyboard.add(button3)
    keyboard = generate_keyboard_requestions(user_id)
    bot.send_message(user_id, "Что включено?", reply_markup=keyboard)


def callback_service(call):
    user_id = call.message.chat.id
    chosen_service = call.data[8:]
    included_services[user_id] = chosen_service



def add_villa(call):
    msg = bot.send_message(call.message.chat.id, "Введите название новой виллы")
    bot.register_next_step_handler(msg, process_add_villa_step)


def process_add_villa_step(message):
    user_id = message.chat.id
    new_villa = message.text
    villa_list[new_villa] = {}
    villa_name[user_id] = new_villa
    villa_list[new_villa]["villa_name"] = new_villa
    msg = bot.send_message(message.chat.id, "Введите локацию виллы")
    bot.register_next_step_handler(msg, process_villa_location, new_villa)

def process_villa_location(message, villa_name):
    user_id = message.chat.id
    villa_list[villa_name]["villa_location"] = message.text
    villa_location[user_id] = message.text
    msg = bot.send_message(user_id, "Введите название виллы как на Google Картах")
    bot.register_next_step_handler(msg, process_villa_google_name, villa_name)

def process_villa_google_name(message, villa_name):
    user_id = message.chat.id
    villa_list[villa_name]["villa_google_name"] = message.text
    villa_google_name[user_id] = message.text
    print_villa_info(user_id,villa_name)

def print_villa_info(user_id, villa_name):
    info = f"Информация о вилле:\n"
    info += f"1. Название виллы: {villa_name}\n"
    info += f"2. Локация виллы: {villa_list[villa_name]['villa_location']}\n"
    info += f"3. Название виллы как на Google Картах: {villa_list[villa_name]['villa_google_name']}\n"
    keyboard = types.InlineKeyboardMarkup()
    button1 = types.InlineKeyboardButton("Заново", callback_data='villas_repeat')
    button2 = types.InlineKeyboardButton("Дальше", callback_data='villas_next')
    button3 = types.InlineKeyboardButton("Сохранить", callback_data='villas_save')
    keyboard.add(button1, button2)
    keyboard.add(button3)
    bot.send_message(user_id, info, reply_markup=keyboard)


def villa_save_callback(message):
    user_id = message.chat.id


def natural_save(call):
    user_id = call.message.chat.id
    user_data = {
        'type': 'natural_person',
        'name': natural_name[user_id],
        'citizenship': natural_citizenship[user_id],
        'passport': natural_passport[user_id],
        'phone': natural_phone[user_id],
        'address': natural_address[user_id],
    }
    msg = bot.send_message(user_id, "Пожалуйста, введите внутреннее имя для этих данных:")
    bot.register_next_step_handler(msg, process_save_name_step, user_data)


def villa_save(call):
    user_id = call.message.chat.id
    user_data = {
        villa_name[user_id]: {
            "villa_name": villa_name[user_id],
            "villa_location": villa_location[user_id],
            "villa_google_name": villa_google_name[user_id]
        }
    }

    msg = bot.send_message(user_id, "Пожалуйста, введите внутреннее имя для этих данных:")
    bot.register_next_step_handler(msg, process_save_name_step_villa, user_data)

def process_save_name_step_villa(message, user_data):
    user_id = message.chat.id
    data_name = message.text

    save_data_villa(user_id, user_data, data_name)
    bot.send_message(user_id, "Ваши данные были сохранены под именем '{}'.".format(data_name))
    choose_villa(message)

def process_save_name_step(message, user_data):
    user_id = message.chat.id
    data_name = message.text

    save_data(user_id, user_data, data_name)
    bot.send_message(user_id, "Ваши данные были сохранены под именем '{}'.".format(data_name))
    print_natural_info(user_id, message)

def legal_save(call):
    user_id = call.message.chat.id

    user_data = {
        'type': 'legal_entity',
        'name': legal_name[user_id],
        'director_name': legal_director_name[user_id],
        'director_citizenship': legal_director_citizenship[user_id],
        'director_passport': legal_director_passport[user_id],
        'nib': legal_nib[user_id],
        'address': legal_address[user_id],
    }
    msg = bot.send_message(user_id, "Пожалуйста, введите внутреннее имя для этих данных:")
    bot.register_next_step_handler(msg, process_save_name_step_legal, user_data)


def process_save_name_step_legal(message, user_data):
    user_id = message.chat.id
    data_name = message.text

    # Сохраняем данные
    save_data(user_id, user_data, data_name)
    bot.send_message(user_id, "Ваши данные были сохранены под именем '{}'.".format(data_name))
    print_company_info(user_id, message)

def handle_long_term(call):
    print(3)

def handle_fsides(call):
    user_id = call.message.chat.id
    data = load_data(user_id)
    bot.edit_message_text(chat_id=call.message.chat.id, message_id=call.message.message_id,
                          text='Выберите первую сторону договора',
                          disable_web_page_preview=True)
    keyboard = types.InlineKeyboardMarkup()
    button1 = types.InlineKeyboardButton("Saba Real Estate", callback_data='f_first_company')
    button2 = types.InlineKeyboardButton("Вторая компания", callback_data='f_second_company')
    button3 = types.InlineKeyboardButton("Третья компания", callback_data='f_third_company')
    button4 = types.InlineKeyboardButton("Человек", callback_data='f_natural_person')
    button5 = types.InlineKeyboardButton("Компания", callback_data='f_legal_entity')
    delete_button = types.InlineKeyboardButton("Удалить", callback_data='delete')
    keyboard.add(button1)
    keyboard.add(button4)
    keyboard.add(button5)

    for person_name in data.keys():
        button = types.InlineKeyboardButton(person_name, callback_data=f'first_{person_name}')
        keyboard.add(button)
    keyboard.add(delete_button)
    bot.edit_message_reply_markup(chat_id=call.message.chat.id, message_id=call.message.message_id,
                                  reply_markup=keyboard)


def ask_for_delete_villa(call):
    msg = bot.send_message(call.message.chat.id, "Введите имя, которое вы хотите удалить.")
    bot.register_next_step_handler(msg, delete_name_villa)

def delete_name_villa(message):
    name_to_delete = message.text
    user_id = message.chat.id
    delete_data_villa(user_id, name_to_delete)
    bot.send_message(user_id, f"'{name_to_delete}' была успешно удалена.")
    choose_villa(message)



def ask_for_delete(call):
    msg = bot.send_message(call.message.chat.id, "Введите имя, которое вы хотите удалить.")
    bot.register_next_step_handler(msg, delete_name,call)


def delete_name(message,call):
    name_to_delete = message.text
    user_id = message.chat.id
    delete_data(user_id, name_to_delete)
    bot.send_message(user_id, f"'{name_to_delete}' был успешно удален.")
    if first_side[user_id] == '':
        data = load_data(user_id)
        keyboard = types.InlineKeyboardMarkup()
        button1 = types.InlineKeyboardButton("Saba Real Estate", callback_data='f_first_company')
        button2 = types.InlineKeyboardButton("Вторая компания", callback_data='f_second_company')
        button3 = types.InlineKeyboardButton("Третья компания", callback_data='f_third_company')
        button4 = types.InlineKeyboardButton("Человек", callback_data='f_natural_person')
        button5 = types.InlineKeyboardButton("Компания", callback_data='f_legal_entity')
        delete_button = types.InlineKeyboardButton("Удалить", callback_data='delete')
        keyboard.add(button1)
        keyboard.add(button4)
        keyboard.add(button5)

        for person_name in data.keys():
            button = types.InlineKeyboardButton(person_name, callback_data=f'first_{person_name}')
            keyboard.add(button)
        keyboard.add(delete_button)
        bot.send_message(chat_id=message.chat.id, text='Выберите первую сторону договора', reply_markup=keyboard,
                         disable_web_page_preview=True)
    else:
        data = load_data(user_id)
        keyboard = types.InlineKeyboardMarkup()
        button1 = types.InlineKeyboardButton("Saba Real Estate", callback_data='s_first_company')
        button2 = types.InlineKeyboardButton("Вторая компания", callback_data='s_second_company')
        button3 = types.InlineKeyboardButton("Третья компания", callback_data='s_third_company')
        button4 = types.InlineKeyboardButton("Человек", callback_data='s_natural_person')
        button5 = types.InlineKeyboardButton("Компания", callback_data='s_legal_entity')
        delete_button = types.InlineKeyboardButton("Удалить", callback_data='delete')
        keyboard.add(button1)
        keyboard.add(button4)
        keyboard.add(button5)
        for person_name in data.keys():
            button = types.InlineKeyboardButton(person_name, callback_data=f'second_{person_name}')
            keyboard.add(button)
        keyboard.add(delete_button)
        bot.send_message(chat_id=message.chat.id, text='Выберите вторую сторону договора', reply_markup=keyboard,
                         disable_web_page_preview=True)




def handle_ssides(call):
    user_id = call.message.chat.id
    data = load_data(user_id)
    bot.edit_message_text(chat_id=call.message.chat.id, message_id=call.message.message_id,
                          text='Выберите вторую сторону договора',
                          disable_web_page_preview=True)
    keyboard = types.InlineKeyboardMarkup()
    button1 = types.InlineKeyboardButton("Saba Real Estate", callback_data='s_first_company')
    button2 = types.InlineKeyboardButton("Вторая компания", callback_data='s_second_company')
    button3 = types.InlineKeyboardButton("Третья компания", callback_data='s_third_company')
    button4 = types.InlineKeyboardButton("Человек", callback_data='s_natural_person')
    button5 = types.InlineKeyboardButton("Компания", callback_data='s_legal_entity')
    delete_button = types.InlineKeyboardButton("Удалить", callback_data='delete')
    keyboard.add(button1)
    keyboard.add(button4)
    keyboard.add(button5)
    for person_name in data.keys():
        button = types.InlineKeyboardButton(person_name, callback_data=f'second_{person_name}')
        keyboard.add(button)
    keyboard.add(delete_button)
    bot.edit_message_reply_markup(chat_id=call.message.chat.id, message_id=call.message.message_id,
                                  reply_markup=keyboard)


def handle_natural_person(call):
    bot.send_message(call.message.chat.id, "Пожалуйста, введите следующую информацию (всю на английском языке):")
    msg = bot.send_message(call.message.chat.id, "1. Фамилия и имя")
    bot.register_next_step_handler(msg, process_name_step)

def process_name_step(message):
    user_id = message.chat.id
    natural_name[user_id] = message.text
    if first_side[user_id] == 'natural_person':
        first_side_name[user_id] = natural_name[user_id]
    else:
        second_side_name[user_id] = natural_name[user_id]
    msg = bot.send_message(user_id, "2. Гражданство")
    bot.register_next_step_handler(msg, process_citizenship_step)

def process_citizenship_step(message):
    user_id = message.chat.id
    natural_citizenship[user_id] = message.text
    msg = bot.send_message(user_id, "3. Паспорт")
    bot.register_next_step_handler(msg, process_passport_step)

def process_passport_step(message):
    user_id = message.chat.id
    natural_passport[user_id] = message.text
    if first_side[user_id] == 'natural_person':
        first_side_passport[user_id] = natural_passport[user_id]
    else:
        second_side_passport[user_id] = natural_passport[user_id]
    msg = bot.send_message(user_id, "4. Номер телефона")
    bot.register_next_step_handler(msg, process_phone_step)

def process_phone_step(message):
    user_id = message.chat.id
    natural_phone[user_id] = message.text
    msg = bot.send_message(user_id, "5. Адрес проживания")
    bot.register_next_step_handler(msg, process_address_step)

def process_address_step(message):
    user_id = message.chat.id
    natural_address[user_id] = message.text
    if first_side[user_id] == 'natural_person':
        first_side_address[user_id] = natural_address[user_id]
    else:
        second_side_address[user_id] = natural_address[user_id]
    msg = bot.send_message(user_id, "6. Прикрепите фотографию паспорта")
    bot.register_next_step_handler(msg, process_passport_photo_step)

def process_passport_photo_step(message):
    user_id = message.chat.id
    passport_photo = message.photo[-1].file_id if message.photo else None
    bot.send_message(user_id, "Спасибо, ваша информация была сохранена.")
    if first_side[user_id] == 'natural_person':
        first_side_name[user_id] = natural_name[user_id]
    else:
        second_side_name[user_id] = natural_name[user_id]
    print_natural_info(user_id,message)

def print_natural_info(user_id, call):
    info = f"Информация о человеке:\n"
    info += f"1. Фамилия и имя: {natural_name[user_id]}\n"
    info += f"2. Гражданство: {natural_citizenship[user_id]}\n"
    info += f"3. Паспорт: {natural_passport[user_id]}\n"
    info += f"4. Номер телефона: {natural_phone[user_id]}\n"
    info += f"5. Адрес проживания: {natural_address[user_id]}\n"
    info += f"6. Фотография паспорта: (сохранена в системе)\n"

    keyboard = types.InlineKeyboardMarkup()
    button1 = types.InlineKeyboardButton("Заново", callback_data='natural_repeat')
    button2 = types.InlineKeyboardButton("Дальше", callback_data='natural_next')
    button3 = types.InlineKeyboardButton("Сохранить", callback_data='natural_save')
    keyboard.add(button1, button2)
    keyboard.add(button3)
    bot.send_message(user_id, info, reply_markup=keyboard)


def save_data_villa(user_id, data, person_name):
    try:
        with open(f'villas_{user_id}.json', 'r') as f:
            all_data = json.load(f)
    except FileNotFoundError:
        all_data = {}

    all_data[person_name] = data
    with open(f'villas_{user_id}.json', 'w') as f:
        json.dump(all_data, f)


def load_data_villa(user_id):
    try:
        with open(f'villas_{user_id}.json', 'r') as f:
            data = json.load(f)
    except FileNotFoundError:
        data = {}
    return data

def delete_data_villa(user_id,person_name):
    try:
        with open(f'villas_{user_id}.json', 'r') as f:
            all_data = json.load(f)
    except FileNotFoundError:
        all_data = {}

    if person_name in all_data:
        del all_data[person_name]

    with open(f'villas_{user_id}.json', 'w') as f:
        json.dump(all_data, f)

def delete_data(user_id, person_name):
    try:
        with open(f'{user_id}.json', 'r') as f:
            all_data = json.load(f)
    except FileNotFoundError:
        all_data = {}

    if person_name in all_data:
        del all_data[person_name]

    with open(f'{user_id}.json', 'w') as f:
        json.dump(all_data, f)


def save_data(user_id, data, person_name):
    try:
        with open(f'{user_id}.json', 'r') as f:
            all_data = json.load(f)
    except FileNotFoundError:
        all_data = {}

    all_data[person_name] = data
    with open(f'{user_id}.json', 'w') as f:
        json.dump(all_data, f)


def load_data(user_id):
    try:
        with open(f'{user_id}.json', 'r') as f:
            data = json.load(f)
    except FileNotFoundError:
        data = {}
    return data

def handle_legal_entity(call):

    bot.send_message(call.message.chat.id, "Пожалуйста, введите следующую информацию:")
    msg = bot.send_message(call.message.chat.id, "1. Название компании")
    bot.register_next_step_handler(msg, process_company_name_step)

def process_company_name_step(message):
    user_id = message.chat.id
    legal_name[user_id] = message.text
    if first_side[user_id] == 'legal_entity':
        first_side_name[user_id] = legal_name[user_id]
    else:
        second_side_name[user_id] = legal_name[user_id]
    msg = bot.send_message(user_id, "2. Фамилия и имя директора")
    bot.register_next_step_handler(msg, process_director_name_step)

def process_director_name_step(message):
    user_id = message.chat.id
    legal_director_name[user_id] = message.text
    msg = bot.send_message(user_id, "3. Гражданство директора")
    bot.register_next_step_handler(msg, process_director_citizenship_step)

def process_director_citizenship_step(message):
    user_id = message.chat.id
    legal_director_citizenship[user_id] = message.text
    msg = bot.send_message(user_id, "4. Паспорт директора")
    bot.register_next_step_handler(msg, process_director_passport_step)

def process_director_passport_step(message):
    user_id = message.chat.id
    legal_director_passport[user_id] = message.text
    msg = bot.send_message(user_id, "5. Прикрепите фотографию паспорта директора")
    bot.register_next_step_handler(msg, process_director_passport_photo_step)

def process_director_passport_photo_step(message):
    user_id = message.chat.id
    legal_director_passport_photo[user_id] = message.photo[-1].file_id if message.photo else None
    msg = bot.send_message(user_id, "6. NIB компании")
    bot.register_next_step_handler(msg, process_nib_step)

def process_nib_step(message):
    user_id = message.chat.id
    legal_nib[user_id] = message.text
    if first_side[user_id] == 'legal_entity':
        first_side_nib[user_id] = legal_nib[user_id]
    else:
        second_side_nib[user_id] = legal_nib[user_id]
    msg = bot.send_message(user_id, "7. Адрес регистрации компании")
    bot.register_next_step_handler(msg, process_company_address_step)

def process_company_address_step(message):
    user_id = message.chat.id
    legal_address[user_id] = message.text
    if first_side[user_id] == 'legal_entity':
        first_side_address[user_id] = legal_address[user_id]
    else:
        second_side_address[user_id] = legal_address[user_id]
    bot.send_message(user_id, "Спасибо, информация о вашей компании была сохранена.")
    print_company_info(user_id, message)


def print_company_info(user_id,call):
    info = f"Информация о компании:\n"
    info += f"1. Название компании: {legal_name[user_id]}\n"
    info += f"2. Фамилия и имя директора: {legal_director_name[user_id]}\n"
    info += f"3. Гражданство директора: {legal_director_citizenship[user_id]}\n"
    info += f"4. Паспорт директора: {legal_director_passport[user_id]}\n"
    info += f"5. Фотография паспорта директора: (сохранена в системе)\n"
    info += f"6. NIB компании: {legal_nib[user_id]}\n"
    info += f"7. Адрес регистрации компании: {legal_address[user_id]}"

    keyboard = types.InlineKeyboardMarkup()
    button1 = types.InlineKeyboardButton("Заново", callback_data='legal_repeat')
    button2 = types.InlineKeyboardButton("Дальше", callback_data='legal_next')
    button3 = types.InlineKeyboardButton("Сохранить", callback_data='legal_save')
    keyboard.add(button1, button2)
    keyboard.add(button3)
    bot.send_message(user_id, info,reply_markup=keyboard)

def generate_keyboard_requestions(user_id):
    keyboard = types.InlineKeyboardMarkup(row_width=1)
    buttons = []
    for service in services:
        text = service
        if service in selected_services.get(user_id, []):
            text += " ✅"
        buttons.append(types.InlineKeyboardButton(text, callback_data=f'service_{service}'))
    keyboard.add(*buttons)
    button8 = types.InlineKeyboardButton('Готово', callback_data='ok')
    keyboard.row(button8)
    return keyboard


def str_to_relativedelta(duration_str):
    quantity, unit = duration_str.split('_')
    quantity = int(quantity)

    if unit == 'month':
        return relativedelta(months=quantity)
    elif unit == 'years':
        return relativedelta(years=quantity)
    else:
        return None


if __name__ == '__main__':
    try:
        bot.infinity_polling(none_stop=True, timeout=90, long_polling_timeout=10)
    except Exception as e:
        print(f"Произошла ошибка: {e}")

